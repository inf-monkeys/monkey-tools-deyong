from flask import Flask, request, jsonify, send_file
import requests
from flask_restx import Api, Resource, fields
import traceback
import logging
import base64
import os
import tempfile
import uuid
import shutil
from docx import Document
import openai
import io
import json
import asyncio
import aiohttp
import re
import time
import urllib.parse
from typing import List, Dict, Any
from lxml import etree
from werkzeug.utils import secure_filename
from datetime import datetime

# 导入腾讯云OCR SDK
# 注意：需要安装 tencentcloud-sdk-python
try:
    from tencentcloud.common import credential
    from tencentcloud.common.profile.client_profile import ClientProfile
    from tencentcloud.common.profile.http_profile import HttpProfile
    from tencentcloud.ocr.v20181119 import ocr_client, models
except ImportError:
    print("请安装腾讯云SDK: pip install tencentcloud-sdk-python")

app = Flask(__name__, static_folder=None)
api = Api(
    app,
    version="1.0",
    title="API Services",
    description="API Services including Document AI Translation and Dify QA",
)

# 导入文件服务所需模块
from flask import send_file

# 文件托管相关配置
class Config:
    # 文件托管目录
    OUTPUT_FILES_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'output_files')
    # 确保目录存在
    os.makedirs(OUTPUT_FILES_DIR, exist_ok=True)
    # 文件访问URL前缀，默认为本地开发环境
    FILE_ACCESS_URL_PREFIX = "http://localhost:5001/files/"
    # 文件过期时间（秒）
    FILE_EXPIRY_SECONDS = 3600 * 24 * 7  # 7天
    
# 尝试从环境变量或配置文件加载配置
try:
    if os.path.exists('config.py'):
        from config import Config as UserConfig
        Config.FILE_ACCESS_URL_PREFIX = getattr(UserConfig, 'FILE_ACCESS_URL_PREFIX', Config.FILE_ACCESS_URL_PREFIX)
        Config.FILE_EXPIRY_SECONDS = getattr(UserConfig, 'FILE_EXPIRY_SECONDS', Config.FILE_EXPIRY_SECONDS)
    elif os.environ.get('FILE_ACCESS_URL_PREFIX'):
        Config.FILE_ACCESS_URL_PREFIX = os.environ.get('FILE_ACCESS_URL_PREFIX')
    
    print(f"文件访问URL前缀: {Config.FILE_ACCESS_URL_PREFIX}")
except Exception as e:
    print(f"加载配置失败: {str(e)}")

# 设置异步翻译的最大并发请求数
MAX_CONCURRENT_REQUESTS = 10
# 每批处理的文本数量
BATCH_SIZE = 20

# 设置API密钥和URL
API_URL = "https://api.cursorai.art"
DIFY_API_URL = "https://api.dify.ai/v1"

# 特定词语的固定翻译
SPECIAL_TRANSLATIONS = {
    "签名": "Ký tên",
    "簽名": "Ký tên",  # 繁体
    "序号": "STT",
    "序號": "STT",  # 繁体
    "会签单位": "Đơn vị ký hiệu",
    "會簽單位": "Đơn vị ký hiệu",  # 繁体
    "编制": "Được soạn bởi",
    "編制": "Được soạn bởi",  # 繁体
    "提议": "Đề nghị",
    "提議": "Đề nghị",  # 繁体
    "审核": "Xét duyệt",
    "審核": "Xét duyệt",  # 繁体
    "核准": "Phê duyệt",
    "復核": "Duyệt lại",  # 繁体
    "复核": "Duyệt lại",
    "VND": "đ",
    "流程": "Quy trình",
    "流程图": "Sơ đồ quy trình",
    "发奖": "Phát thưởng",
    "情形": "Tình hình",
    "处理": "Xử lý",
    "开单": "Mở phiếu",
    "说明": "Giải thích",
    "受奖人": "Người nhận thưởng",
    "签字": "Ký tên",
    "批准": "Phê duyệt",
    "公告": "Thông báo",
    "生效": "Có hiệu lực",
    "存档": "Lưu trữ",
    "系统": "Hệ thống",
    "档案": "Hồ sơ",
    # 数字相关翻译
    "百": "00",
    "佰": "00",  # 繁体
    "千": "000",
    "仟": "000",  # 繁体
    "万": "0000",
    "萬": "0000",  # 繁体
    # 添加表格中常见的单元格标题
    "奖励种类": "Loại khen thưởng",
    "奖勰種類": "Loại khen thưởng",  # 繁体
    "组织": "Tổ chức",
    "組織": "Tổ chức",  # 繁体
    "部门": "Bộ phận",
    "部門": "Bộ phận",  # 繁体
    "公司": "Công ty",
    "人事部": "Phòng nhân sự"
}

# 创建上传文件的目录
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 限制上传文件大小为16MB
app.config['OUTPUT_FILES_DIR'] = Config.OUTPUT_FILES_DIR

ai_translation_ns = api.namespace("ai_translation", description="Document Translation API")
ocr_ns = api.namespace("ocr", description="腾讯云OCR API")
dify_ns = api.namespace("dify", description="Dify API")
inference_ns = api.namespace("inference", description="Inference API")

# Define document translation request and response models
document_translation_request = ai_translation_ns.model(
    "DocumentTranslationRequest",
    {
        "url": fields.String(required=True, description="URL of the document to translate"),
        "target_language": fields.String(required=True, description="Target language for translation"),
        "special_requirements": fields.String(required=False, description="Special requirements for translation"),
        "api_key": fields.String(required=False, description="API key for translation service")
    }
)

document_translation_response = ai_translation_ns.model(
    "DocumentTranslationResponse",
    {
        "message": fields.String(description="Response message"),
        "document_url": fields.String(description="URL to the translated document"),
        "error": fields.String(description="Error message if any")
    }
)

class NoSuccessfulRequestLoggingFilter(logging.Filter):
    def filter(self, record):
        return "GET /" not in record.getMessage()

# 获取 Flask 的默认日志记录器
log = logging.getLogger("werkzeug")
# 创建并添加过滤器
log.addFilter(NoSuccessfulRequestLoggingFilter())

@app.before_request
def before_request():
    request.app_id = request.headers.get("x-monkeys-appid")
    request.user_id = request.headers.get("x-monkeys-userid")
    request.team_id = request.headers.get("x-monkeys-teamid")
    request.workflow_id = request.headers.get("x-monkeys-workflowid")
    request.workflow_instance_id = request.headers.get("x-monkeys-workflow-instanceid")

# 添加静态文件托管路由
@app.route('/files/<path:filename>')
def serve_file(filename):
    """提供对文件的访问"""
    return send_file(os.path.join(Config.OUTPUT_FILES_DIR, filename))

@api.errorhandler(Exception)
def handle_exception(error):
    return {"message": str(error)}, 500

@app.get("/manifest.json")
def get_manifest():
    return {
        "schema_version": "v1",
        "display_name": "Deyong",
        "namespace": "monkey_tools_deyong",
        "auth": {"type": "none"},
        "api": {"type": "openapi", "url": "/swagger.json"},
        "contact_email": "dev@inf-monkeys.com",
        "categories": ["ai", "translation", "document"],
        "description": {
            "zh-CN": "Deyong Tools",
            "en-US": "Deyong Tools"
        },
        "icon": "emoji:📄:#3a8fe5",
        "credentials": [
            {
                "name": "cursor-ai",
                "type": "aksk",
                "displayName": {
                    "zh-CN": "Cursor AI",
                    "en-US": "Cursor AI"
                },
                "iconUrl": "emoji:🤖:#3a8fe5",
                "properties": [
                    {
                        "displayName": {
                            "zh-CN": "从 Cursor AI 获取你的 API Key",
                            "en-US": "Get your API Key from Cursor AI"
                        },
                        "type": "notice",
                        "name": "docs"
                    },
                    {
                        "displayName": {
                            "zh-CN": "API Key",
                            "en-US": "API Key"
                        },
                        "type": "string",
                        "name": "api_key",
                        "required": True
                    }
                ]
            },
            {
                "name": "dify",
                "type": "aksk",
                "displayName": {
                    "zh-CN": "Dify",
                    "en-US": "Dify"
                },
                "iconUrl": "https://dify.ai/favicon.ico",
                "properties": [
                    {
                        "displayName": {
                            "zh-CN": "从 [Dify](https://dify.ai) 获取你的 API Key",
                            "en-US": "Get your API Key from [Dify](https://dify.ai)"
                        },
                        "type": "notice",
                        "name": "docs"
                    },
                    {
                        "displayName": {
                            "zh-CN": "API Key",
                            "en-US": "API Key"
                        },
                        "type": "string",
                        "name": "api_key",
                        "required": True
                    }
                ]
            }
        ]
    }

@app.route('/upload', methods=['POST'])
def upload_file():
    file_base64 = None
    file_type = None
    
    # 检查是否是从本地上传的文件
    if 'file' in request.files and request.files['file'].filename != '':
        file = request.files['file']
        
        # 获取文件类型
        file_type = file.filename.rsplit('.', 1)[1].lower() if '.' in file.filename else ''
        if file_type not in ['jpg', 'jpeg', 'png', 'bmp', 'pdf']:
            return jsonify({'error': '不支持的文件类型'}), 400
        
        # 保存文件
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        
        # 将文件转换为base64编码
        with open(file_path, 'rb') as f:
            file_content = f.read()
            file_base64 = base64.b64encode(file_content).decode('utf-8')
        
        # 删除临时文件
        os.remove(file_path)
        
    # 检查是否提供了CDN URL
    elif 'cdn_url' in request.form and request.form['cdn_url'] != '':
        cdn_url = request.form['cdn_url']
        
        try:
            # 从URL获取文件内容
            response = requests.get(cdn_url, stream=True)
            response.raise_for_status() # 确保请求成功
            
            # 从URL中提取文件类型
            content_type = response.headers.get('Content-Type', '')
            if 'image/jpeg' in content_type:
                file_type = 'jpg'
            elif 'image/png' in content_type:
                file_type = 'png'
            elif 'image/bmp' in content_type:
                file_type = 'bmp'
            elif 'application/pdf' in content_type:
                file_type = 'pdf'
            else:
                # 尝试从URL中获取文件扩展名
                url_path = urllib.parse.urlparse(cdn_url).path
                file_type = url_path.rsplit('.', 1)[1].lower() if '.' in url_path else ''
            
            if file_type not in ['jpg', 'jpeg', 'png', 'bmp', 'pdf']:
                return jsonify({'error': '不支持的文件类型'}), 400
            
            # 将文件内容转换为base64
            file_content = response.content
            file_base64 = base64.b64encode(file_content).decode('utf-8')
            
        except requests.exceptions.RequestException as e:
            return jsonify({'error': f'无法从CDN下载文件: {str(e)}'}), 400
    else:
        return jsonify({'error': '没有文件被上传或提供CDN URL'}), 400
    
    return jsonify({
        'file_base64': file_base64,
        'file_type': file_type
    })

@ai_translation_ns.route("/document")
class DocumentTranslationResource(Resource):
    @ai_translation_ns.doc("translate_document")
    @ai_translation_ns.vendor(
        {
            "x-monkey-tool-name": "translate_document",
            "x-monkey-tool-categories": ["ai", "translation", "document"],
            "x-monkey-tool-display-name": {
                "zh-CN": "文档AI翻译",
                "en-US": "Document AI Translation",
            },
            "x-monkey-tool-description": {
                "zh-CN": "使用GPT-4o进行文档翻译",
                "en-US": "Document translation using GPT-4o",
            },
            "x-monkey-tool-icon": "emoji:📄:#3a8fe5",
            "x-monkey-tool-input": [
                {
                    "displayName": {
                        "zh-CN": "Cursor AI API密钥",
                        "en-US": "Cursor AI API Key",
                    },
                    "name": "api_key",
                    "type": "string",
                    "required": True,
                },
                {
                    "displayName": {
                        "zh-CN": "文档CDN URL",
                        "en-US": "Document CDN URL",
                    },
                    "name": "document_url",
                    "type": "string",
                    "required": True,
                },
                {
                    "displayName": {
                        "zh-CN": "目标语言",
                        "en-US": "Target Language",
                    },
                    "name": "target_language",
                    "type": "string",
                    "required": True,
                },
                {
                    "displayName": {
                        "zh-CN": "特殊翻译要求",
                        "en-US": "Special Translation Requirements",
                    },
                    "name": "special_requirements",
                    "type": "string",
                    "required": False,
                }
            ],
            "x-monkey-tool-output": [
                {
                    "displayName": {
                        "zh-CN": "翻译后的文档",
                        "en-US": "Translated Document",
                    },
                    "name": "translated_document",
                    "type": "file",
                }
            ],
            "x-monkey-tool-extra": {
                "estimateTime": 180,
            },
        }
    )
    @ai_translation_ns.expect(document_translation_request)
    @ai_translation_ns.response(200, "成功", document_translation_response)
    @ai_translation_ns.response(400, "请求无效", document_translation_response)
    @ai_translation_ns.response(401, "未授权", document_translation_response)
    @ai_translation_ns.response(500, "服务器错误", document_translation_response)
    def post(self):
        """
        Translate a Word document using GPT-4o
        
        This endpoint accepts a Word document from a CDN URL,
        translates it to the specified target language, and returns a bilingual document 
        with both the original text and the translation.
        
        Returns a Word document with the translated content.
        """
        try:
            # 获取JSON请求数据
            json_data = request.json
            if not json_data:
                return {
                    "file_url": "",
                    "success": False,
                    "message": "无效的请求数据。必须提供有效的JSON数据。"
                }, 400
                
            api_key = json_data.get('api_key')
            if not api_key:
                return {
                    "file_url": "",
                    "success": False,
                    "message": "Missing API key"
                }, 401
                
            target_language = json_data.get('target_language')
            special_requirements = json_data.get('special_requirements', '')
            document_url = json_data.get('document_url')
            
            if not target_language:
                return {
                    "file_url": "",
                    "success": False,
                    "message": "Missing target language parameter"
                }, 400
                
            if not document_url:
                return {
                    "file_url": "",
                    "success": False,
                    "message": "未提供文档CDN URL"
                }, 400
                
            # Create a temporary file to store the document
            temp_dir = tempfile.mkdtemp()
            input_file_path = os.path.join(temp_dir, f"input_{uuid.uuid4()}.docx")
            output_file_path = os.path.join(temp_dir, f"output_{uuid.uuid4()}.docx")
            
            # 从URL中提取文件名
            url_path = urllib.parse.urlparse(document_url).path
            file_name = os.path.basename(url_path)
            if not file_name.endswith('.docx'):
                return {
                    "file_url": "",
                    "success": False,
                    "message": "只支持 .docx 格式的文件"
                }, 400
                
            try:
                # 从URL下载文件
                response = requests.get(document_url, stream=True)
                response.raise_for_status()  # 确保请求成功
                
                # 保存下载的文件
                with open(input_file_path, 'wb') as f:
                    for chunk in response.iter_content(chunk_size=8192):
                        f.write(chunk)
                        
            except requests.exceptions.RequestException as e:
                return {
                    "file_url": "",
                    "success": False,
                    "message": f"无法从CDN URL下载文件: {str(e)}"
                }, 400
                
            # 处理文档
            translated_doc = self.translate_document(input_file_path, target_language, special_requirements, api_key)
            translated_doc.save(output_file_path)
            
            # 删除输入临时文件，但保留输出文件以供上传到S3
            os.remove(input_file_path)
            
            # 创建一个持久化的输出目录，确保S3上传工具能访问到
            output_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'output_files')
            os.makedirs(output_dir, exist_ok=True)
            
            # 生成一个有意义的文件名，包含时间戳和原始文件名
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            original_filename = os.path.basename(url_path)
            filename_base, _ = os.path.splitext(original_filename)
            persistent_filename = f"{filename_base}_{target_language}_{timestamp}.docx"
            persistent_filepath = os.path.join(output_dir, persistent_filename)
            
            # 复制翻译后的文件到持久化目录
            shutil.copy2(output_file_path, persistent_filepath)
            
            # 清理剩余的临时文件
            os.remove(output_file_path)
            os.rmdir(temp_dir)
                
            # 生成可访问的URL
            file_url = f"{Config.FILE_ACCESS_URL_PREFIX}{persistent_filename}"
            
            # 返回文件URL和相关信息
            return {
                "file_path": persistent_filepath,     # 本地文件系统路径（用于调试）
                "file_url": file_url,                # 可访问的URL
                "publicAccessUrl": file_url,         # 给S3用的公开访问URL
                "filename": persistent_filename,      # 文件名
                "success": True,
                "message": f"文档翻译成功，可通过 {file_url} 访问"
            }
                
        except Exception as e:
            traceback.print_exc()
            return {
                "file_url": "",
                "success": False,
                "message": str(e)
            }, 500

    async def translate_text_async(self, text, session, target_language, special_requirements="", api_key=None):
        """
        使用 GPT-4o API 异步翻译中文文本
        
        Args:
            text: 要翻译的文本
            session: aiohttp 客户端会话
            target_language: 目标语言
            special_requirements: 特殊翻译要求
        
        Returns:
            翻译后的文本
        """
        if not text.strip():
            return ""
        
        # 检查是否为单独的字符或阿拉伯数字
        if len(text.strip()) <= 1 or text.strip().isdigit():
            return text
        
        # 检查是否为特定词语
        # if text.strip() in SPECIAL_TRANSLATIONS:
        #     return SPECIAL_TRANSLATIONS[text.strip()]
        
        try:
            # 构建 API 请求
            headers = {
                "Content-Type": "application/json",
                "Authorization": f"Bearer {api_key}"
            }
            
            # 新的OpenAI API格式要求有user参数
            data = {
                "model": "gpt-4o",
                "messages": [
                    {"role": "system", "content": f"你是一个专业的中文到{target_language}翻译器。请将用户提供的中文文本翻译成{target_language}，只输出翻译结果，不要有任何解释或额外内容。保持原始格式，但不要重复原文中的标点符号，特别是在行尾的标点符号。如果原文中有标点符号，请使用{target_language}中的对应标点符号，而不是重复使用原文的标点符号。如果遇到单独的字母或数字，请保持原样不翻译。如果文本中包含“百”、“千”、“万”等数字单位，请按照特定规则翻译。{special_requirements if special_requirements else ''}"},
                    {"role": "user", "content": text}
                ],
                "temperature": 0.3,
                "user": "translation_service"  # 添加user参数以满足API要求
            }
            
            # 发送 API 请求
            print(f"正在发送翻译请求: {text[:30]}...")
            async with session.post(f"{API_URL}/v1/chat/completions", headers=headers, json=data) as response:
                response_data = await response.json()
                
                # 处理 API 响应
                if response.status == 200 and "choices" in response_data:
                    translated_text = response_data["choices"][0]["message"]["content"]
                    print(f"翻译成功: {translated_text[:30]}...")
                    return translated_text
                else:
                    print(f"翻译失败: {response.status} - {response_data}")
                    return ""
        except Exception as e:
            print(f"翻译过程中发生错误: {str(e)}")
            return ""

    async def batch_translate_texts(self, texts, target_language, special_requirements="", api_key=None):
        """
        批量异步翻译多个文本
        
        Args:
            texts: 要翻译的文本列表
            target_language: 目标语言
            special_requirements: 特殊翻译要求
        
        Returns:
            翻译后的文本列表
        """
        # 创建异步会话
        async with aiohttp.ClientSession() as session:
            # 创建信号量限制并发请求数
            semaphore = asyncio.Semaphore(MAX_CONCURRENT_REQUESTS)
            
            async def translate_with_semaphore(text):
                async with semaphore:
                    return await self.translate_text_async(text, session, target_language, special_requirements, api_key)
            
            # 创建所有翻译任务
            tasks = [translate_with_semaphore(text) for text in texts]
            
            # 等待所有任务完成
            results = await asyncio.gather(*tasks)
            return results
    
    def translate_text(self, text, target_language, special_requirements="", api_key=None):
        """
        同步版本的翻译函数，用于兼容现有代码
        
        Args:
            text: 要翻译的文本
            target_language: 目标语言
            special_requirements: 特殊翻译要求
        
        Returns:
            翻译后的文本
        """
        if not text.strip():
            return ""
        
        # 使用同步方式调用异步函数
        try:
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            result = loop.run_until_complete(self.batch_translate_texts([text], target_language, special_requirements, api_key))[0]
            loop.close()
            return result
        except Exception as e:
            print(f"同步翻译过程中发生错误: {str(e)}")
            return ""
    
    def process_docx(self, input_file_path, target_language, special_requirements, api_key=None):
        """
        处理Word文档，翻译其中的文本并创建双语文档
        
        Args:
            input_file_path: Word文档路径
            target_language: 目标语言
            special_requirements: 特殊翻译要求
            
        Returns:
            翻译后的Document对象
        """
        # 打开原始文档
        doc = Document(input_file_path)
        
        # 翻译正文段落
        total_paragraphs = len(doc.paragraphs)
        print(f"文档共有 {total_paragraphs} 个段落")
        
        # 收集需要翻译的段落文本
        paragraph_texts = []
        paragraph_refs = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                paragraph_texts.append(paragraph.text)
                paragraph_refs.append(paragraph)
        
        # 批量并行翻译段落
        if paragraph_texts:
            print(f"开始批量翻译 {len(paragraph_texts)} 个段落...")
            # 使用异步方式批量翻译
            translated_texts = asyncio.run(self.batch_translate_texts(paragraph_texts, target_language, special_requirements, api_key))
            
            # 创建段落和翻译结果的映射
            paragraphs_to_translate = []
            for paragraph, translated_text in zip(paragraph_refs, translated_texts):
                if translated_text.strip():
                    paragraphs_to_translate.append((paragraph, translated_text))
                else:
                    print(f"  警告: 段落翻译失败，不添加翻译")
        else:
            paragraphs_to_translate = []
        
        # 现在在原文后面添加翻译文本
        # 从后往前遍历，这样我们在添加新段落时不会影响前面的段落索引
        for paragraph, translated_text in reversed(paragraphs_to_translate):
            try:
                # 使用更安全的方式插入翻译文本
                # 直接在段落后面添加一个新段落
                p = doc.add_paragraph()
                # 获取原段落的父元素
                parent_element = paragraph._p.getparent()
                # 获取原段落在父元素中的索引
                if parent_element is not None:
                    index_in_parent = list(parent_element).index(paragraph._p)
                    # 在原段落后面插入新段落
                    parent_element.insert(index_in_parent + 1, p._p)
            except Exception as e:
                print(f"  警告: 插入段落时出错: {str(e)}")
                # 如果插入失败，尝试直接在文档末尾添加段落
                p = doc.add_paragraph()
            
            # 设置翻译文本和样式
            run = p.add_run(translated_text)
            
            # 复制原段落的样式
            if paragraph.style:
                p.style = paragraph.style
            
            # 复制原段落的对齐方式
            if paragraph.alignment is not None:
                p.alignment = paragraph.alignment
            
            # 如果原段落有格式，复制字体格式
            if paragraph.runs:
                # 获取所有格式属性
                for orig_run in paragraph.runs:
                    if orig_run.font.size:
                        run.font.size = orig_run.font.size
                    if orig_run.font.name:
                        run.font.name = orig_run.font.name
                    # 复制加粗、斜体、下划线等格式
                    if hasattr(orig_run.font, 'bold') and orig_run.font.bold:
                        run.font.bold = orig_run.font.bold
                    if hasattr(orig_run.font, 'italic') and orig_run.font.italic:
                        run.font.italic = orig_run.font.italic
                    if hasattr(orig_run.font, 'underline') and orig_run.font.underline:
                        run.font.underline = orig_run.font.underline
                    # 复制颜色
                    if hasattr(orig_run.font, 'color') and orig_run.font.color and hasattr(orig_run.font.color, 'rgb') and orig_run.font.color.rgb:
                        run.font.color.rgb = orig_run.font.color.rgb
                    # 一旦找到有格式的run，就使用它的格式
                    if any([orig_run.font.bold, orig_run.font.italic, orig_run.font.underline, orig_run.font.size]):
                        break
        
        # 处理表格
        all_table_cells = []
        all_table_texts = []
        
        # 收集所有表格单元格的文本
        for table in doc.tables:
            print("正在处理表格...")
            
            for row in table.rows:
                for cell in row.cells:
                    # 获取单元格的文本
                    cell_text = cell.text.strip()
                    
                    if cell_text:
                        all_table_cells.append(cell)
                        all_table_texts.append(cell_text)
        
        # 批量并行翻译表格单元格
        if all_table_texts:
            print(f"开始批量翻译 {len(all_table_texts)} 个表格单元格...")
            # 使用异步方式批量翻译
            translated_table_texts = asyncio.run(self.batch_translate_texts(all_table_texts, target_language, special_requirements, api_key))
            
            # 处理翻译结果
            cell_translations = []
            for cell, translated_text in zip(all_table_cells, translated_table_texts):
                if translated_text.strip():
                    cell_translations.append((cell, translated_text))
                else:
                    print(f"  警告: 表格单元格翻译失败，不添加翻译")
            
            # 创建一个集合来跟踪已处理的单元格，防止重复处理
            processed_cells = set()
            
            # 将翻译结果添加到表格单元格中
            for cell, translated_text in zip(all_table_cells, translated_table_texts):
                # 使用单元格对象的ID作为唯一标识符
                cell_id = id(cell)
                
                # 如果这个单元格已经处理过，则跳过
                if cell_id in processed_cells:
                    continue
                    
                # 标记这个单元格为已处理
                processed_cells.add(cell_id)
                
                try:
                    # 检查单元格是否已经包含翻译
                    already_translated = False
                    
                    # 获取所有段落文本，检查是否已包含翻译
                    all_cell_text = cell.text
                    if translated_text.strip() in all_cell_text:
                        print(f"  跳过已翻译的单元格内容")
                        continue
                    
                    # 逐段检查是否已包含翻译
                    for para in cell.paragraphs[1:] if len(cell.paragraphs) > 1 else []:
                        if para.text.strip() == translated_text.strip():
                            already_translated = True
                            break
                            
                    if already_translated:
                        continue
                    
                    # 添加翻译段落
                    if len(cell.paragraphs) > 0 and cell.paragraphs[0].text.strip():
                        # 添加新段落
                        p = cell.add_paragraph()
                        p.text = translated_text
                        
                        # 尝试应用原始段落的样式
                        if cell.paragraphs[0].style:
                            p.style = cell.paragraphs[0].style
                except Exception as e:
                    print(f"  处理表格单元格时出错: {str(e)}")
        
        return doc
    
    def call_translation_api(self, input_file_path, target_language,api_key):
        """
        调用app.py中的/api/translate接口来翻译文档
        
        Args:
            input_file_path: 输入文档路径
            target_language: 目标语言
            
        Returns:
            翻译后的文档路径
        """
        import requests
        import tempfile
        import os
        
        # 创建临时文件来保存翻译后的文档
        temp_dir = tempfile.mkdtemp()
        output_path = os.path.join(temp_dir, f"translated_output.docx")
        
        # 准备API请求
        url = "http://localhost:5005/api/translate"  # app.py运行的地址
        
        # 准备文件和表单数据
        files = {
            'file': (os.path.basename(input_file_path), open(input_file_path, 'rb'), 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        }
        data = {
            'target_language': target_language,
            'api_key': api_key
        }
        
        try:
            # 发送请求
            print(f"正在调用翻译API...")
            response = requests.post(url, files=files, data=data, stream=True)
            
            # 检查响应
            if response.status_code == 200:
                # 将响应内容保存到文件
                with open(output_path, 'wb') as f:
                    for chunk in response.iter_content(chunk_size=8192):
                        f.write(chunk)
                print(f"翻译成功，结果保存到: {output_path}")
                return output_path
            else:
                print(f"翻译API调用失败: {response.status_code} - {response.text}")
                raise Exception(f"翻译API调用失败: {response.status_code}")
        except Exception as e:
            print(f"调用翻译API时出错: {str(e)}")
            raise e
        finally:
            # 关闭文件
            files['file'][1].close()
    
    def translate_document(self, input_file_path, target_language, special_requirements,api_key):
        """
        Translate a Word document using GPT-4o
        
        Args:
            input_file_path: Path to the input Word document
            target_language: Target language for translation
            special_requirements: Special translation requirements
            
        Returns:
            A Document object with the translated content
        """
        try:
            # 调用翻译API
            output_path = self.call_translation_api(input_file_path, target_language,api_key)
            
            # 返回翻译后的文档
            return Document(output_path)
        except Exception as e:
            print(f"翻译文档时出错: {str(e)}")
            # 如果API调用失败，回退到使用本地翻译方法
            print("尝试使用本地翻译方法...")
            return self.process_docx(input_file_path, target_language, special_requirements, api_key)


# 定义OCR请求模型
ocr_request = ocr_ns.model(
    "OCRRequest",
    {
        "image_url": fields.String(required=True, description="图片的URL地址"),
        "secret_id": fields.String(required=True, description="腾讯云SecretId"),
        "secret_key": fields.String(required=True, description="腾讯云SecretKey"),
    },
)

# 定义响应模型
ocr_response = ocr_ns.model(
    "OCRResponse",
    {
        "extracted_text": fields.String(description="OCR提取的原始文本"),
        "success": fields.Boolean(description="OCR识别是否成功"),
        "message": fields.String(description="处理结果信息")
    },
)

@ocr_ns.route("/extract")
class OCRExtractResource(Resource):
    @ocr_ns.doc("extract_text_from_document")
    @ocr_ns.vendor(
        {
            "x-monkey-tool-name": "extract_text_from_document",
            "x-monkey-tool-categories": ["ocr", "document-processing"],
            "x-monkey-tool-display-name": {
                "zh-CN": "从图片提取文本",
                "en-US": "Extract Text from Image",
            },
            "x-monkey-tool-description": {
                "zh-CN": "使用腾讯云OCR从图片URL中提取文本",
                "en-US": "Extract text from image URL using Tencent Cloud OCR",
            },
            "x-monkey-tool-icon": "emoji:📔:#4a90e2",
            "x-monkey-tool-input": [
                {
                    "displayName": {
                        "zh-CN": "图片URL",
                        "en-US": "Image URL",
                    },
                    "name": "image_url",
                    "type": "string",
                    "required": True,
                },
                {
                    "displayName": {
                        "zh-CN": "腾讯云SecretId",
                        "en-US": "Tencent Cloud SecretId",
                    },
                    "name": "secret_id",
                    "type": "string",
                    "required": True,
                },
                {
                    "displayName": {
                        "zh-CN": "腾讯云SecretKey",
                        "en-US": "Tencent Cloud SecretKey",
                    },
                    "name": "secret_key",
                    "type": "string",
                    "required": True,
                }
            ],
            "x-monkey-tool-output": [
                {
                    "displayName": {
                        "zh-CN": "提取的文本",
                        "en-US": "Extracted Text",
                    },
                    "name": "extracted_text",
                    "type": "string",
                },
                {
                    "displayName": {
                        "zh-CN": "是否成功",
                        "en-US": "Success",
                    },
                    "name": "success",
                    "type": "boolean",
                },
                {
                    "displayName": {
                        "zh-CN": "结果信息",
                        "en-US": "Message",
                    },
                    "name": "message",
                    "type": "string",
                }
            ],
            "x-monkey-tool-extra": {
                "estimateTime": 5,
            },
        }
    )
    @ocr_ns.expect(ocr_request)
    @ocr_ns.response(200, "成功", ocr_response)
    def post(self):
        """
        使用腾讯云OCR从图片URL中提取文本
        """
        json_data = request.json
        image_url = json_data.get("image_url")
        secret_id = json_data.get("secret_id")
        secret_key = json_data.get("secret_key")
        
        # 使用腾讯云OCR提取文本
        extracted_text = self.perform_ocr_from_url(image_url, secret_id, secret_key)
        
        if extracted_text.startswith("OCR错误"):
            return {
                "extracted_text": "",
                "success": False,
                "message": extracted_text
            }
        
        return {
            "extracted_text": extracted_text,
            "success": True,
            "message": "文本提取成功"
        }
    
    def perform_ocr_from_url(self, image_url, secret_id, secret_key):
        """使用腾讯云OCR API从图片URL提取文本"""
        try:
            # 创建认证对象
            cred = credential.Credential(secret_id, secret_key)
            
            # 创建客户端配置
            httpProfile = HttpProfile()
            httpProfile.endpoint = "ocr.tencentcloudapi.com"  # API网关地址
            httpProfile.reqMethod = "POST"  # 请求方法
            httpProfile.reqTimeout = 30    # 超时时间，单位为秒
            
            clientProfile = ClientProfile()
            clientProfile.httpProfile = httpProfile
            clientProfile.signMethod = "TC3-HMAC-SHA256"  # 签名方法
            
            # 创建OCR客户端，默认使用广州区域
            client = ocr_client.OcrClient(cred, "ap-guangzhou", clientProfile)
            
            # 创建请求对象
            req = models.GeneralBasicOCRRequest()
            
            # 设置图片URL
            req.ImageUrl = image_url
            
            # 可选参数设置
            # req.LanguageType = "auto"  # 识别语言类型，默认为自动
            # req.Scene = "normal"       # 场景值，默认为通用
            # req.IsWords = False        # 是否返回单字信息
            
            # 调用通用印刷体识别接口
            response = client.GeneralBasicOCR(req)
            
            # 提取文本和位置信息
            result = []
            text_items = []
            for item in response.TextDetections:
                text_items.append(item.DetectedText)
                result.append({
                    "text": item.DetectedText,  # 识别出的文本
                    "confidence": item.Confidence,  # 置信度
                    "polygon": {  # 文本框坐标
                        "x": [item.Polygon[0].X, item.Polygon[1].X, item.Polygon[2].X, item.Polygon[3].X],
                        "y": [item.Polygon[0].Y, item.Polygon[1].Y, item.Polygon[2].Y, item.Polygon[3].Y]
                    } if hasattr(item, 'Polygon') and item.Polygon else None
                })
            
            # 打印详细结果信息（调试用）
            print(f"OCR识别结果: {response.to_json_string()}")
            
            # 返回纯文本结果
            return "\n".join(text_items)
        
        except Exception as e:
            print(f"OCR错误: {str(e)}")
            return f"OCR错误: {str(e)}"


# 定义Dify QA请求模型
dify_request = dify_ns.model(
    "DifyRequest",
    {
        "api_key": fields.String(required=True, description="Dify API密钥"),
        "question": fields.String(required=True, description="要提问的问题"),
        "conversation_id": fields.String(required=False, description="对话ID，用于继续之前的对话"),
    },
)

# 定义Dify QA响应模型
dify_response = dify_ns.model(
    "DifyResponse",
    {
        "answer": fields.String(description="AI回答的内容"),
        "conversation_id": fields.String(description="对话ID"),
        "success": fields.Boolean(description="请求是否成功")
    },
)

# 定义文档翻译请求模型
document_translation_request = ai_translation_ns.model(
    "DocumentTranslationRequest",
    {
        "document_url": fields.String(required=True, description="文档CDN URL，必须是.docx格式文件"),
        "api_key": fields.String(required=True, description="Cursor AI API密钥"),
        "target_language": fields.String(required=True, description="目标翻译语言"),
        "special_requirements": fields.String(required=False, description="特殊翻译要求")
    },
)

# 定义文档翻译响应模型（虽然实际响应是一个文件）
document_translation_response = ai_translation_ns.model(
    "DocumentTranslationResponse",
    {
        "file_url": fields.String(description="翻译后的文档URL"),
        "success": fields.Boolean(description="翻译是否成功"),
        "message": fields.String(description="处理结果信息")
    },
)
@dify_ns.route("/qa")
class DifyQAResource(Resource):
    @dify_ns.doc("qa_service")
    @dify_ns.expect(dify_request)
    @dify_ns.response(200, "成功", dify_response)
    @dify_ns.vendor(
        {
            "x-monkey-tool-name": "dify_qa",
            "x-monkey-tool-categories": ["ai", "qa"],
            "x-monkey-tool-display-name": {
                "zh-CN": "Dify问答服务",
                "en-US": "Dify QA Service",
            },
            "x-monkey-tool-description": {
                "zh-CN": "使用Dify API进行问答",
                "en-US": "Use Dify API for QA",
            },
            "x-monkey-tool-icon": "emoji:📄:#3a8fe5",
            "x-monkey-tool-input": [
                {
                    "displayName": {
                        "zh-CN": "Dify API密钥",
                        "en-US": "Dify API Key",
                    },
                    "name": "api_key",
                    "type": "string",
                    "required": True,
                },
                {
                    "displayName": {
                        "zh-CN": "问题",
                        "en-US": "Question",
                    },
                    "name": "question",
                    "type": "string",
                    "required": True,
                },
                {
                    "displayName": {
                        "zh-CN": "对话ID",
                        "en-US": "Conversation ID",
                    },
                    "name": "conversation_id",
                    "type": "string",
                    "required": False,
                }
            ],
            "x-monkey-tool-output": [
                {
                    "displayName": {
                        "zh-CN": "回答",
                        "en-US": "Answer",
                    },
                    "name": "answer",
                    "type": "string",
                },
                {
                    "displayName": {
                        "zh-CN": "对话ID",
                        "en-US": "Conversation ID",
                    },
                    "name": "conversation_id",
                    "type": "string",
                },
                {
                    "displayName": {
                        "zh-CN": "成功",
                        "en-US": "Success",
                    },
                    "name": "success",
                    "type": "boolean",
                }
            ],
            "x-monkey-tool-extra": {
                "estimateTime": 5,
            },
        }
    )
    def post(self):
        
            # 获取请求数据
            data = request.json
            api_key = data.get("api_key")
            if not api_key:
                return {"error": "Missing Dify API key"}, 401
                
            question = data.get("question")
            conversation_id = data.get("conversation_id", "")


            if not question:
                return {"error": "问题不能为空"}, 400
            
            # 准备请求头
            headers = {
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json"
            }
            user_id = "user-" + str(hash(datetime.now().strftime('%Y%m%d%H%M%S')))
            # 准备请求数据
            data = {
                "inputs": {},
                "query": question,
                "user": user_id,
                "response_mode": "blocking",
            }

            # 仅当会话ID存在且有效时才添加到请求中
            import re
            uuid_pattern = re.compile(r'^[0-9a-f]{8}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{12}$', re.I)
            if conversation_id and (isinstance(conversation_id, str) and uuid_pattern.match(conversation_id)):
                data["conversation_id"] = conversation_id
            
            try:
                # 发送请求到Dify API
                response = requests.post(
                    f"{DIFY_API_URL}/chat-messages",
                    headers=headers,
                    json=data
                )
                
                if response.status_code == 200:
                    result = response.json()
                    answer = result.get("answer", "抱歉，我无法回答这个问题。")
                    
                    # 返回结果
                    return {
                        "answer": answer,
                        "conversation_id": result.get("conversation_id", ""),
                        "success": True
                    }
                else:
                    return {"error": f"API请求失败: {response.text}"}, response.status_code
            
            except Exception as e:
                error_msg = f"发生错误: {str(e)}"
                return {"answer": error_msg, "success": False}


def extract_formulas_from_response(response_text: str) -> List[str]:
    """
    从 GPT-o3 API 的响应文本中提取数学公式
    
    Args:
        response_text: GPT-o3 返回的文本响应
        
    Returns:
        提取出的公式列表
    """
    formulas = []
    
    # 尝试查找常见的公式标记
    formula_markers = [
        "公式：", "公式:", "公式是", "公式为", "表达式：", "表达式:", 
        "数学公式：", "数学公式:", "formula:", "formula：", "equation:", "equation：",
        "f(x) =", "f(n) =", "y =", "Y =", "output =", "a_n ="
    ]
    
    # 分割文本为行
    lines = response_text.split('\n')
    
    # 遍历每一行查找公式
    for line in lines:
        line = line.strip()
        
        # 跳过空行
        if not line:
            continue
            
        # 检查是否包含公式标记
        for marker in formula_markers:
            if marker in line:
                # 提取公式部分
                formula_part = line[line.find(marker):].strip()
                if formula_part and len(formula_part) > len(marker):
                    formulas.append(formula_part)
                    break
        
        # 检查是否有 Markdown 代码块中的公式
        if line.startswith('```') and ('math' in line or 'latex' in line):
            # 查找代码块结束
            in_code_block = True
            code_block_content = []
            for next_line in lines[lines.index(line) + 1:]:  
                if next_line.strip() == '```':
                    in_code_block = False
                    break
                code_block_content.append(next_line.strip())
            
            if code_block_content and not in_code_block:
                formulas.append(''.join(code_block_content))
    
    # 如果没有找到明确的公式标记，尝试查找可能的公式模式
    if not formulas:
        # 查找包含等号和数学符号的行
        math_symbols = ['+', '-', '*', '/', '^', '=', '(', ')', '[', ']', '{', '}', '\\', 'sqrt', 'log', 'sin', 'cos']
        for line in lines:
            line = line.strip()
            if '=' in line and any(sym in line for sym in math_symbols):
                formulas.append(line)
    
    # 去重并返回
    return list(set(formulas))


def analyze_data_patterns(data_points: List[float]) -> Dict[str, Any]:
    """
    分析数据点之间的规律和关系
    """
    results = {}
    
    # 基本统计信息
    results["count"] = len(data_points)
    results["min"] = min(data_points)
    results["max"] = max(data_points)
    results["mean"] = np.mean(data_points)
    results["median"] = np.median(data_points)
    
    # 检查等差数列
    differences = [data_points[i+1] - data_points[i] for i in range(len(data_points)-1)]
    if len(set(round(diff, 6) for diff in differences)) == 1:
        results["arithmetic_sequence"] = True
        results["common_difference"] = differences[0]
        results["formula"] = f"a_n = {data_points[0]} + (n-1) * {differences[0]}"
    else:
        results["arithmetic_sequence"] = False
    
    # 检查等比数列
    if all(x > 0 for x in data_points):
        ratios = [data_points[i+1] / data_points[i] for i in range(len(data_points)-1)]
        if len(set(round(ratio, 6) for ratio in ratios)) == 1:
            results["geometric_sequence"] = True
            results["common_ratio"] = ratios[0]
            results["formula"] = f"a_n = {data_points[0]} * ({ratios[0]})^(n-1)"
        else:
            results["geometric_sequence"] = False
    else:
        results["geometric_sequence"] = False
    
    # 检查二次函数关系
    if len(data_points) >= 3:
        x = np.array(range(1, len(data_points) + 1))
        y = np.array(data_points)
        
        # 线性拟合
        linear_coeffs = np.polyfit(x, y, 1)
        linear_y_pred = np.polyval(linear_coeffs, x)
        linear_residuals = y - linear_y_pred
        linear_mse = np.mean(linear_residuals ** 2)
        
        # 二次拟合
        quad_coeffs = np.polyfit(x, y, 2)
        quad_y_pred = np.polyval(quad_coeffs, x)
        quad_residuals = y - quad_y_pred
        quad_mse = np.mean(quad_residuals ** 2)
        
        # 找出最佳拟合
        if linear_mse < 1e-6:
            results["best_fit"] = "linear"
            results["formula"] = f"f(n) = {linear_coeffs[0]:.6f}*n + {linear_coeffs[1]:.6f}"
            results["coefficients"] = linear_coeffs.tolist()
        elif quad_mse < 1e-6:
            results["best_fit"] = "quadratic"
            results["formula"] = f"f(n) = {quad_coeffs[0]:.6f}*n^2 + {quad_coeffs[1]:.6f}*n + {quad_coeffs[2]:.6f}"
            results["coefficients"] = quad_coeffs.tolist()
    
    return results

def call_gpt_o3(json_data, api_key) -> Dict[str, Any]:
    """
    调用 GPT-o3 API 来进行数据推理
    """
    try:
        # 构建 API 请求
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {api_key}"
        }
        
        # 构建示例数据
        # example_data = {
        #   "data": [
        #     {
        #       "instruction": "X1X2X3:193.0,22.0,88.0",
        #       "input": "",
        #       "output": "Y1Y2Y3:42.18,65.42,8.9"
        #     },
        #     {
        #       "instruction": "X1X2X3:243.0,175.0,76.0",
        #       "input": "",
        #       "output": "Y1Y2Y3:76.22,15.6,58.75"
        #     }
        #   ]
        # }
        example_data = {
          "data": [
            {
              
            }
          ]
        }
        
        
        # 构建GPT-o3的提示词
        prompt = """以下是JSON格式的数据，instruction和output之间存在关联。请完成以下任务：
1. 分析instruction和output之间的数学关系
2. 给出能够从instruction推导出output的精确公式
3. 创建一个完整的Python代码文件，该文件应包含：
   - 一个函数，能够接收instruction格式的输入并返回对应的output
   - 清晰的注释，解释数据之间的关系和转换逻辑
   - 示例代码，展示如何使用该函数处理示例数据
   - 必要的辅助函数和数据处理逻辑

请确保Python代码是完整的、可执行的，并能准确表达数据之间的关系。
    """
        
        # 如果用户提供了自定义数据，使用用户数据，否则使用示例数据
        data_to_analyze = json_data if json_data else example_data
        # data_to_analyze = example_data
        user_message = f"{prompt}\n\n{json.dumps(data_to_analyze, ensure_ascii=False, indent=2)}"
        
        messages = [
            {"role": "system", "content": "你是一个数据分析专家，擅长发现数据之间的规律和公式。"},
            {"role": "user", "content": user_message}
        ]
        
        data = {
            "model": "gpt-4o",  # 指定使用 GPT-o3 模型 ，目前用api跑o3有点问题，先暂时用4o
            "messages": messages,
            "temperature": 0.3  # 设置温度
        }
        
        # 发送 API 请求
        print(f"正在发送数据分析请求...")
        response = requests.post(f"{API_URL}/v1/chat/completions", headers=headers, json=data, timeout=30)
        
        # 处理 API 响应
        if response.status_code == 200 and "choices" in response.json():
            gpt_response = response.json()
            ai_message = gpt_response["choices"][0]["message"]["content"]
            print(f"分析成功!")
            
            # 从响应中提取Python代码
            python_code = ""
            # 查找Python代码块
            code_blocks = re.findall(r'```python\n([\s\S]*?)```', ai_message)
            if code_blocks:
                python_code = code_blocks[0]
            
            # 返回结果
            return {
                "input_data": json_data,
                "analysis_time": time.time(),
                "prompt_used": prompt,
                "analysis_result": ai_message,
                "python_code": python_code
            }
        else:
            # API 调用失败
            error_msg = f"分析失败: {response.status_code} - {response.text}"
            print(error_msg)
            return {
                "input_data": json_data,
                "error": error_msg
            }
    except Exception as e:
        # 异常处理
        error_msg = f"API 调用异常: {str(e)}"
        print(error_msg)
        traceback.print_exc()
        return {
            "input_data": json_data,
            "error": error_msg
        }


@inference_ns.route("/o3")
class InferenceO3Resource(Resource):
    @inference_ns.doc("infer_data_patterns")
    @inference_ns.vendor(
        {
            "x-monkey-tool-name": "infer_data_patterns",
            "x-monkey-tool-categories": ["data_analysis", "inference"],
            "x-monkey-tool-display-name": {
                "zh-CN": "数据规律推理",
                "en-US": "Data Pattern Inference",
            },
            "x-monkey-tool-description": {
                "zh-CN": "分析数据之间的规律和可换算的公式",
                "en-US": "Analyze patterns and formulas between data points",
            },
            "x-monkey-tool-icon": "emoji:📊:#4a90e2",
            "x-monkey-tool-input": [
                {
                    "displayName": {
                        "zh-CN": "Cursor AI API密钥",
                        "en-US": "Cursor AI API Key",
                    },
                    "name": "api_key",
                    "type": "string",
                    "required": True,
                },
                {
                    "displayName": {
                        "zh-CN": "数据点",
                        "en-US": "Data Points",
                    },
                    "name": "data_points",
                    "type": "array",
                    "required": True,
                    "description": {
                        "zh-CN": "要分析的数据点列表",
                        "en-US": "List of data points to analyze",
                    }
                },
                {
                    "displayName": {
                        "zh-CN": "分析模式",
                        "en-US": "Analysis Mode",
                    },
                    "name": "analysis_mode",
                    "type": "string",
                    "required": False,
                    "description": {
                        "zh-CN": "分析模式，支持 'basic' 和 'advanced'",
                        "en-US": "Analysis mode, supports 'basic' and 'advanced'",
                    }
                }
            ],
            "x-monkey-tool-output": [
                {
                    "displayName": {
                        "zh-CN": "数据点",
                        "en-US": "Data Points",
                    },
                    "name": "data_points",
                    "type": "array",
                },
                {
                    "displayName": {
                        "zh-CN": "分析时间",
                        "en-US": "Analysis Time",
                    },
                    "name": "analysis_time",
                    "type": "number",
                },
                {
                    "displayName": {
                        "zh-CN": "使用的提示词",
                        "en-US": "Prompt Used",
                    },
                    "name": "prompt_used",
                    "type": "string",
                },
                {
                    "displayName": {
                        "zh-CN": "基本规律",
                        "en-US": "Basic Patterns",
                    },
                    "name": "basic_patterns",
                    "type": "object",
                },
                {
                    "displayName": {
                        "zh-CN": "高级洞察",
                        "en-US": "Advanced Insights",
                    },
                    "name": "advanced_insights",
                    "type": "object",
                }
            ],
            "x-monkey-tool-extra": {
                "estimateTime": 10,
                "provider": "GPT-o3",
            },
            "x-monkey-tool-credentials": [
                {
                    "name": "cursor-ai",
                    "required": True,
                    "description": {
                        "zh-CN": "Cursor AI API 密钥",
                        "en-US": "Cursor AI API Key"
                    }
                }
            ]
        }
    )
    @inference_ns.expect(
        inference_ns.model(
            "DataInferenceRequest",
            {
                "api_key": fields.String(required=True, description="Cursor AI API密钥"),
                "data": fields.Raw(description="Any valid JSON data, including arrays and objects")
            }
        )
    )
    @inference_ns.response(
        200,
        "Success",
        inference_ns.model(
            "DataInferenceResult",
            {
                "input_data": fields.Raw(description="The input JSON data"),
                "analysis_time": fields.Float(description="Analysis timestamp"),
                "prompt_used": fields.String(description="The prompt used for GPT-o3 analysis"),
                "analysis_result": fields.String(description="Analysis result from GPT-o3"),
                "python_code": fields.String(description="Extracted Python code that represents the data relationship")
            },
        ),
    )
    def post(self):
        """
        分析数据点之间的规律和可换算的公式
        """
        try:
            # 获取请求数据
            request_data = request.json
            if request_data is None:
                return {"message": "Invalid request data. Must provide valid JSON data."}, 400
            
            # 获取API密钥
            api_key = request_data.get('api_key')
            if not api_key:
                return {"error": "Missing API key"}, 401
                
            # 如果请求中有 data 字段，则使用该字段的值
            # 否则直接使用整个请求数据
            json_data = request_data.get('data', request_data)
            
            # 直接将 JSON 数据发送给 GPT-o3 进行分析
            result = call_gpt_o3(json_data, api_key)
            
            return result
            
        except Exception as e:
            # 异常处理
            traceback.print_exc()
            return {"message": f"Error analyzing data: {str(e)}"}, 500


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5001)
